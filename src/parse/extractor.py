"""Извлечение текста из тендерной документации.

Поддерживаемые форматы:
- PDF (родной): pymupdf
- PDF (скан): pymupdf + tesseract — fallback (тесьма не подключаем по умолчанию)
- DOCX: python-docx
- DOC: антиквариатный формат — fallback через `antiword`/`textract` (не входит в deps)
- RTF: striprtf или поток через regex
- XLSX/XLS: openpyxl/xlrd — отдельный pipeline для таблиц тех. характеристик

Возврат — `ExtractedDocument` со структурой:
    text: полный текст
    paragraphs: список абзацев с грубыми позициями
    metadata: формат, число страниц, время извлечения
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

log = logging.getLogger(__name__)

WHITESPACE_RE = re.compile(r"[ \t ]+")
MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


@dataclass
class ExtractedDocument:
    file_path: Path
    file_format: str
    text: str
    paragraphs: list[str] = field(default_factory=list)
    n_pages: int | None = None
    extraction_method: str = ""
    error: str | None = None


def _normalize(text: str) -> str:
    text = WHITESPACE_RE.sub(" ", text)
    text = MULTI_NEWLINE_RE.sub("\n\n", text)
    # Decoy soft-hyphens, NBSPs, zero-widths
    text = text.replace("­", "").replace("​", "")
    return text.strip()


def extract_pdf(path: Path) -> ExtractedDocument:
    import pymupdf  # type: ignore

    pages: list[str] = []
    n = 0
    with pymupdf.open(path) as doc:
        n = doc.page_count
        for page in doc:
            txt = page.get_text("text")
            pages.append(txt)
    raw = "\n\n".join(pages)
    text = _normalize(raw)
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    return ExtractedDocument(
        file_path=path,
        file_format="pdf",
        text=text,
        paragraphs=paras,
        n_pages=n,
        extraction_method="pymupdf",
    )


def extract_docx(path: Path) -> ExtractedDocument:
    import docx  # python-docx

    d = docx.Document(str(path))
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    # Tables — flatten cells row-by-row
    for tbl in d.tables:
        for row in tbl.rows:
            row_text = " | ".join(
                _normalize(cell.text) for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paras.append(row_text)
    text = _normalize("\n\n".join(paras))
    return ExtractedDocument(
        file_path=path,
        file_format="docx",
        text=text,
        paragraphs=[p for p in text.split("\n\n") if p.strip()],
        extraction_method="python-docx",
    )


def extract_rtf(path: Path) -> ExtractedDocument:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Поставь `pip install striprtf` для RTF") from exc
    raw = path.read_text(encoding="utf-8", errors="ignore")
    text = _normalize(rtf_to_text(raw))
    return ExtractedDocument(
        file_path=path,
        file_format="rtf",
        text=text,
        paragraphs=[p for p in text.split("\n\n") if p.strip()],
        extraction_method="striprtf",
    )


def extract_doc_legacy(path: Path) -> ExtractedDocument:
    """Старый .doc — пробуем антиквариат.
    Возвращает пустой текст если не получилось (без падения)."""
    import subprocess

    methods = [
        (["antiword", "-w", "0", str(path)], "antiword"),
        (["catdoc", str(path)], "catdoc"),
    ]
    last_err = ""
    for cmd, name in methods:
        try:
            r = subprocess.run(cmd, capture_output=True, timeout=30)
            if r.returncode == 0 and r.stdout:
                text = _normalize(r.stdout.decode("utf-8", errors="ignore"))
                return ExtractedDocument(
                    file_path=path,
                    file_format="doc",
                    text=text,
                    paragraphs=[p for p in text.split("\n\n") if p.strip()],
                    extraction_method=name,
                )
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            last_err = str(e)
    return ExtractedDocument(
        file_path=path,
        file_format="doc",
        text="",
        extraction_method="failed",
        error=f"no .doc extractor available: {last_err}",
    )


def extract_xlsx(path: Path) -> ExtractedDocument:
    """Excel — собираем текст по ячейкам, по строкам, по листам.
    Часто ТЗ/НМЦК хранят в xlsx, текст там значимый."""
    import openpyxl  # type: ignore

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    paras: list[str] = []
    for sheet in wb.worksheets:
        paras.append(f"=== Лист: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c not in (None, "")]
            if cells:
                paras.append(" | ".join(cells))
    text = _normalize("\n".join(paras))
    return ExtractedDocument(
        file_path=path,
        file_format="xlsx",
        text=text,
        paragraphs=[p for p in text.split("\n") if p.strip()],
        extraction_method="openpyxl",
    )


def extract_xls(path: Path) -> ExtractedDocument:
    """Старый Excel — через xlrd (только формат .xls 97-2003)."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        return ExtractedDocument(
            file_path=path, file_format="xls", text="",
            extraction_method="missing-deps", error="xlrd not installed",
        )
    book = xlrd.open_workbook(str(path))
    paras: list[str] = []
    for sheet in book.sheets():
        paras.append(f"=== Лист: {sheet.name} ===")
        for row_i in range(sheet.nrows):
            cells = [str(sheet.cell_value(row_i, c)) for c in range(sheet.ncols)
                    if sheet.cell_value(row_i, c) not in ("", None)]
            if cells:
                paras.append(" | ".join(cells))
    text = _normalize("\n".join(paras))
    return ExtractedDocument(
        file_path=path, file_format="xls", text=text,
        paragraphs=[p for p in text.split("\n") if p.strip()],
        extraction_method="xlrd",
    )


def extract_odt(path: Path) -> ExtractedDocument:
    """OpenDocument Text — внутри это zip с content.xml."""
    import zipfile
    import re

    try:
        with zipfile.ZipFile(path) as z:
            with z.open("content.xml") as f:
                xml = f.read().decode("utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        return ExtractedDocument(
            file_path=path, file_format="odt", text="",
            extraction_method="error", error=str(exc),
        )
    # Грубо вытаскиваем текст из xml — заменяем теги переходами и собираем
    text = re.sub(r"<text:p[^>]*>", "\n", xml)
    text = re.sub(r"<text:h[^>]*>", "\n# ", text)
    text = re.sub(r"<text:tab/>", "\t", text)
    text = re.sub(r"<text:line-break/>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)  # убираем все остальные xml-теги
    # Декодируем HTML entities
    text = (text.replace("&amp;", "&").replace("&lt;", "<")
            .replace("&gt;", ">").replace("&quot;", '"').replace("&apos;", "'"))
    text = _normalize(text)
    return ExtractedDocument(
        file_path=path, file_format="odt", text=text,
        paragraphs=[p for p in text.split("\n\n") if p.strip()],
        extraction_method="zipfile+regex",
    )


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".rtf": extract_rtf,
    ".doc": extract_doc_legacy,
    ".xlsx": extract_xlsx,
    ".xls": extract_xls,
    ".odt": extract_odt,
}


def extract(path: Path) -> ExtractedDocument:
    """Универсальная точка входа: dispatch по расширению."""
    suffix = path.suffix.lower()
    fn = EXTRACTORS.get(suffix)
    if not fn:
        return ExtractedDocument(
            file_path=path,
            file_format=suffix.lstrip("."),
            text="",
            extraction_method="unsupported",
            error=f"unsupported format: {suffix}",
        )
    try:
        return fn(path)
    except Exception as exc:  # noqa: BLE001
        log.warning("extract failed for %s: %s", path, exc)
        return ExtractedDocument(
            file_path=path,
            file_format=suffix.lstrip("."),
            text="",
            extraction_method="error",
            error=str(exc),
        )
